"""
Skapar Strategiportfoljen_Beskrivning_v310.docx
- Ny titelsida: vit bakgrund med blå kurva-bild, svart text
- Baserad på v303, versionssträngar uppdaterade
- Ny avslutningssida: Nyheter i version 3.10
"""
import os
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DIR   = r"C:\Users\hejma\Projekt_Claude\strategiportfoljen"
src   = fr"{DIR}\Strategiportfoljen_Beskrivning_v303.docx"
# Om _ny.docx finns (öppen fil-situation), använd den
if not os.path.exists(src):
    src = fr"{DIR}\Strategiportfoljen_Beskrivning_v303_ny.docx"
dst   = fr"{DIR}\Strategiportfoljen_Beskrivning_v310.docx"
IMG   = fr"{DIR}\titelbild_v310.png"

NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0x0E, 0xA5, 0xE9)
MUTED  = RGBColor(0x6B, 0x72, 0x80)
BLACK  = RGBColor(0x11, 0x18, 0x27)

src_doc = Document(src)

# Hitta första sidbryttet
pb_idx = None
for i, para in enumerate(src_doc.paragraphs):
    for br in para._element.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page":
            pb_idx = i
            break
    if pb_idx is not None:
        break

new_doc = Document()
for sec in new_doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

# ── NY TITELSIDA: bild + minimal text ────────────────────────────
new_doc.add_picture(IMG, width=Inches(6.1))
p = new_doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)

# Sidbyte
new_doc.add_paragraph().add_run().add_break(
    __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE
)

# ── KOPIERA INNEHÅLLSSIDOR FRÅN v303 ─────────────────────────────
src_children = list(src_doc.element.body)
new_body     = new_doc.element.body

for child in src_children[pb_idx + 1:]:
    new_body.append(deepcopy(child))

# ── ERSÄTT 3.03 → 3.10 ───────────────────────────────────────────
for para in new_doc.paragraphs:
    for run in para.runs:
        for old, new in [("v3.03","v3.10"),("v303","v310"),
                         ("version 3.03","version 3.10"),
                         ("Version 3.03","Version 3.10"),
                         ("3.03","3.10")]:
            if old in run.text:
                run.text = run.text.replace(old, new)
for tbl in new_doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if "3.03" in run.text:
                        run.text = run.text.replace("3.03","3.10")

# ── NY SIDA: NYHETER I VERSION 3.10 ──────────────────────────────
def sidbyte(doc):
    doc.add_paragraph().add_run().add_break(
        __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)

def rubrik1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

def rubrik2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ACCENT

def brödtext(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = BLACK

def punkt(doc, text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    if prefix:
        r1 = p.add_run(prefix + " "); r1.bold = True
        r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
        r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK
    else:
        r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = BLACK

sidbyte(new_doc)
rubrik1(new_doc, "9. Nyheter i version 3.10")
rubrik2(new_doc, "Dashboard, kontoregister, importguide och avstämning")
brödtext(new_doc,
    "Version 3.10 är en stor uppgradering med fokus på visualisering, tillförlitlighet "
    "och pedagogisk tydlighet — appen ska vara lika lätt att förstå som att använda."
)

rubrik2(new_doc, "Nytt portföljutvecklingsdiagram")
brödtext(new_doc,
    "Längst upp på Dashboard visas nu ett interaktivt diagram med egna periodknappar "
    "(default: i År). Välj fritt vilka dataserier du vill jämföra via kryssrutor."
)
punkt(new_doc, "Portföljvärde, Nettoinsatt kapital, Nettoresultat som valbara linjer.")
punkt(new_doc, "Per-kategori-linjer — aktivera enskilda kategorier för att jämföra utveckling.")
punkt(new_doc, "Växla mellan linjegraf och stapeldiagram.")

rubrik2(new_doc, "Kontoregister — auto-synk vid namnbyte")
brödtext(new_doc,
    "Appen lagrar nu ett register över kontonummer → kontonamn. Om du byter namn på "
    "ett konto i Avanza uppdateras appen automatiskt vid nästa positionsimport. "
    "Tidigare namn sparas som historik och visas med 🔁 i Avstämning."
)

rubrik2(new_doc, "Importordningsguide")
brödtext(new_doc,
    "Importera-fliken visar en steg-för-steg-guide med live-status (✅/⚠️/⬜) för de "
    "4 importstegen: Transaktioner → Positioner → Inköpskurser → Excel-backup. "
    "Guiden förhindrar att inköpskurser importeras innan positioner finns."
)

rubrik2(new_doc, "Förbättrad diff-wizard i Avstämning")
brödtext(new_doc,
    "Steg 2 förklarar nu exakt var siffrorna finns i Avanza "
    "(Min ekonomi → Sparande → 'Totalt sparande'). "
    "Steg 3 visar en prioriterad beslutsgraf i stället för en platt lista."
)
punkt(new_doc, "✅ Allt stämmer (diff < 200 kr) — ingen åtgärd.")
punkt(new_doc, "⚠️ Importera ny positionsfil — vanligaste orsaken till diff.")
punkt(new_doc, "Hämta FX-kurser om USD/EUR-innehav visar fel värde.")
punkt(new_doc, "Kassadiff — uppdateras automatiskt vid positionsimport.")

rubrik2(new_doc, "Förbättrade nyckeltalskort")
brödtext(new_doc,
    "Nyckeltalskorten på Dashboard har fått ikoner (💼 📈 💳 📊 🏦 🌱 ✅), "
    "bättre ordning och en förklarande rad per kort som tydliggör vad varje mått mäter."
)

p = new_doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Strategiportföljen v3.10  ·  Byggt för Martin  ·  Strategi från januari 2026")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED

new_doc.save(dst)
print(f"Skapad: {dst}")
