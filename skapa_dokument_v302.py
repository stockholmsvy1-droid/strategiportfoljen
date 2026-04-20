"""
Skapar Strategiportfoljen_Beskrivning_v302.docx
Baserad på v301 — uppdaterad med v3.02-nyheter
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, shutil, os

# Kopiera v301 som bas och öppna den
src = r"C:\Users\hejma\Projekt_Claude\strategiportfoljen\Strategiportfoljen_Beskrivning_v301.docx"
dst = r"C:\Users\hejma\Projekt_Claude\strategiportfoljen\Strategiportfoljen_Beskrivning_v302.docx"
shutil.copy2(src, dst)

doc = Document(dst)

NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0x0E, 0xA5, 0xE9)
MUTED  = RGBColor(0x6B, 0x72, 0x80)
BLACK  = RGBColor(0x11, 0x18, 0x27)

# ── Ersätt versionsnummer i alla paragrafer ──
for para in doc.paragraphs:
    for run in para.runs:
        if 'v3.01' in run.text or 'v301' in run.text or 'version 3.01' in run.text or '3.01' in run.text:
            run.text = run.text.replace('v3.01','v3.02').replace('v301','v302').replace('version 3.01','version 3.02').replace('3.01','3.02')
        if 'Version 3.01' in run.text:
            run.text = run.text.replace('Version 3.01','Version 3.02')

# ── Ersätt i tabeller ──
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if '3.01' in run.text:
                        run.text = run.text.replace('3.01','3.02')

# ── Lägg till ny sida med v3.02-nyheter ──
def sidbyte(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(__import__('docx.enum.text', fromlist=['WD_BREAK']).WD_BREAK.PAGE)

def rubrik1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(16); run.font.color.rgb = NAVY
    return p

def rubrik2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(12); run.font.color.rgb = ACCENT
    return p

def brödtext(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5); run.font.color.rgb = BLACK
    return p

def punkt(doc, text, prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    if prefix:
        r1 = p.add_run(prefix + " "); r1.bold = True
        r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
        r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK
    else:
        run = p.add_run(text); run.font.size = Pt(10.5); run.font.color.rgb = BLACK
    return p

sidbyte(doc)

rubrik1(doc, "7. Nyheter i version 3.02")
rubrik2(doc, "Avstämning, periodbaserat nettoinsatt och förbättrad Kassa-flik")

brödtext(doc,
    "Version 3.02 fokuserar på kontroll och konsistens — det ska vara lätt att "
    "verifiera att appen räknar rätt och att siffrorna stämmer mot Avanza."
)

rubrik2(doc, "Ny flik: Avstämning")
brödtext(doc,
    "En helt ny flik låter dig stämma av appens beräknade portföljvärde mot vad "
    "Avanza visar i Min ekonomi. Fliken byggs automatiskt från den senast importerade "
    "positionsfilen — ingen manuell inmatning krävs."
)
punkt(doc, "Saldon per konto — kontonamn (exakt som i Avanza), tillgängligt för köp och totalt marknadsvärde per konto.")
punkt(doc, "Kontroll mot app — differens i kr och % mellan positionsfilens summavärde och appens beräknade portföljvärde.")
punkt(doc, "Förklaringstext — inbyggd förklaring av varför en liten differens är normal (FX-kursskillnader, manuell kassa, dolda konton).")
brödtext(doc,
    "Kontona visas i samma ordning som i Avanza och uppdateras automatiskt vid varje positionsimport. "
    "En differens under 1 % är acceptabel och beror normalt på att appen och Avanza använder "
    "något olika valutakurser för utländska innehav."
)

rubrik2(doc, "Nettoinsatt kapital — nu periodbaserat")
brödtext(doc,
    "Nyckeltalskortet 'Nettoinsatt kapital' på Dashboard följer nu vald period. "
    "Med filtret 'i År' visas nettoinflödet under innevarande år, med totalt nettoinsatt "
    "sedan start som undertext. Vid 'Allt' visas det ackumulerade totalt som tidigare."
)

rubrik2(doc, "Kassa-fliken — Avanza-ordning")
brödtext(doc,
    "Konton i 'Tillgängligt för köp'-tabellen och väljarmenyn visas nu i samma ordning "
    "som Avanza visar dem i Min ekonomi, inte längre alfabetiskt."
)

# Sidfot
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Strategiportföljen v3.02  ·  Byggt för Martin  ·  Strategi från januari 2026")
run.italic = True; run.font.size = Pt(9); run.font.color.rgb = MUTED

doc.save(dst)
print(f"Skapad: {dst}")
