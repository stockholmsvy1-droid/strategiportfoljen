"""
Skapar Strategiportfoljen_Beskrivning_v313.docx
Fristående dokument — ej kopierat från tidigare version.
Texttitelssida (ingen bild). Utförlig beskrivning av v3.13.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DIR = r"C:\Users\hejma\Projekt_Claude\strategiportfoljen"
dst = fr"{DIR}\Strategiportfoljen_Beskrivning_v313.docx"

NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0x0E, 0xA5, 0xE9)
MUTED  = RGBColor(0x6B, 0x72, 0x80)
BLACK  = RGBColor(0x11, 0x18, 0x27)
RED    = RGBColor(0xDC, 0x26, 0x26)

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

WD_BREAK = __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK

def sidbyte():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def tom_rad(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)

def titel_huvud(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY

def titel_version(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(16); r.font.color.rgb = ACCENT

def titel_tagline(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(12); r.font.color.rgb = MUTED

def rubrik1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

def rubrik2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ACCENT

def ingress(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = NAVY

def brödtext(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.color.rgb = BLACK

def punkt(text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    if prefix:
        r1 = p.add_run(prefix + "  "); r1.bold = True
        r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
        r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK
    else:
        r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = BLACK

def citat(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent  = Cm(1.5)
    r = p.add_run(f'"{text}"')
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = MUTED

def fotnot():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("Strategiportföljen v3.13  ·  Byggt för Martin  ·  Strategi från januari 2026")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED


# ══════════════════════════════════════════════
#  TITELSSIDA
# ══════════════════════════════════════════════
for _ in range(6): tom_rad(8)
titel_huvud("Strategiportföljen")
titel_version("version 3.13")
titel_tagline("The Magnificent Martin's Money-Making Machine")
tom_rad(4)
brödtext(
    "En personlig portföljapp för att följa en aktieportfölj med fokus på "
    "långsiktig tillväxt, riskmedvetenhet och strategisk disciplin."
)
fotnot()


# ══════════════════════════════════════════════
#  1. VAD ÄR STRATEGIPORTFÖLJEN?
# ══════════════════════════════════════════════
sidbyte()
rubrik1("1. Vad är Strategiportföljen?")
ingress("v3.13: Stabil grund — Avstämning mot Avanza fungerar nu exakt som Avanzas egna översikt")

brödtext(
    "Strategiportföljen är en personlig webbapp byggd för att ge en samlad, överskådlig bild av "
    "en aktieportfölj som följer en genomtänkt investeringsstrategi med sex kategorier. "
    "Appen hjälper dig att hålla disciplinen — att följa strategin i stället för att reagera på brus."
)
brödtext(
    "Tanken bakom appen är enkel: du ska kunna öppna den på måndag morgon, importera en CSV-fil "
    "från Avanza, och direkt se om något kräver din uppmärksamhet — en MA200-signal, ett "
    "ombalanseringsbehov, en nödutgång. Allt annat är tyst."
)

rubrik2("Arkitektur och filosofi")
brödtext(
    "Appen är avsiktligt byggd som en enda HTML-fil — index.html — utan backend, utan databas, "
    "utan inloggning. All data sparas i webbläsarens localStorage och kan exporteras som Excel-backup. "
    "Inga prenumerationer, inga servrar, inga tredjepartsberoenden utöver Chart.js och SheetJS."
)
brödtext(
    "Det är ett medvetet val som speglar investeringsfilosofin: enkelhet minskar fel. Precis som "
    "en portfölj med för många innehav blir svår att förvalta, blir en app med för många beroenden "
    "svår att underhålla och lita på."
)

rubrik2("Tre användningssätt")
punkt("Veckovis uppföljning — Importera positionsfil från Avanza. Granska nyckeltal, MA200-signaler och Avstämning mot Avanza.", "→")
punkt("Beslutsunderlag — Innan ett köp eller sälj: verifiera kategoritillhörighet, MA200-avstånd och portföljvikt.", "→")
punkt("Strategisk reflektion — Beslutsloggen och värdeutvecklingsdiagrammet ger en historisk bild av vad som fungerat.", "→")


# ══════════════════════════════════════════════
#  2. INVESTERINGSSTRATEGIN
# ══════════════════════════════════════════════
sidbyte()
rubrik1("2. Investeringsstrategin bakom appen")
ingress("Sex kategorier, tydliga roller — strukturen är allt")

brödtext(
    "Appen är designad kring en specifik investeringsstrategi med sex kategorier. "
    "Varje kategori representerar en roll i portföljen — en uppgift den ska utföra, "
    "inte bara en bransch den tillhör. Det är skillnaden mellan en portfölj och en samling aktier."
)

rubrik2("De sex kategorierna")
punkt(
    "Basen i portföljen. Köps regelbundet, säljs aldrig vid dipp. Ger exponering mot bred "
    "marknadsuppgång med minimal risk och låg kostnad. Avanza Zero, Länsförsäkringar Global, Spiltan.",
    "⚓ Ankaret  (Indexfonder · 35–40 %)"
)
punkt(
    "Stabila bolag med lång utdelningshistorik. Utdelningarna återinvesteras. "
    "Säljs inte vid marknadskorrigeringar — de är portföljens kassaflöde. Handelsbanken, Nordea, Johnson & Johnson.",
    "💰 Kassaflödet  (Utdelningsaktier · 20–25 %)"
)
punkt(
    "NVIDIA, TSMC, ASML, Micron — bolag som bygger fundamentet för nästa teknologivåg. "
    "MA200-regeln med tvådagarsbekräftelse gäller. Volatila men strukturellt viktiga.",
    "🔵 Infrastrukturen  (AI-hårdvara / chip · 15–20 %)"
)
punkt(
    "Bolag som tjänar pengar på att AI faktiskt används — Alphabet, Palantir. "
    "Högre värdering men exponering mot en sekulär tillväxttrend.",
    "🧠 Hjärnan  (AI-mjukvara · 10–15 %)"
)
punkt(
    "SAAB, BAE Systems, Kongsberg, Leonardo. Defensiv exponering mot ett strukturellt "
    "underinvesterat segment. MA200-reglerna gäller.",
    "🛡️ Skölden  (Försvarsindustri · 5–10 %)"
)
punkt(
    "Spekulativa positioner med hög conviction. Liten andel, stor potentiell uppsida — "
    "och full risk för förlust. International Petroleum, Klarna. Behandlas som lotter: begränsad storlek.",
    "✨ Berättelserna  (Kryddor / teman · 0–5 %)"
)

rubrik2("MA200-regeln — din disciplinerade exit")
brödtext(
    "För kategorierna 3–6 gäller en tydlig regel: om en aktie stänger under sitt 200-dagars "
    "glidande medelvärde (MA200) i lokal valuta i två dagar i rad, är det en säljsignal. "
    "Inte en rekommendation — en regel. Regeln eliminerar det svåraste investeringsbeslutet: "
    "när ska jag sälja?"
)
brödtext(
    "MA200 anges alltid i lokal valuta — NVIDIA i USD, SAAB i SEK. Det eliminerar bruset från "
    "SEK-fluktuationer som annars ger falska signaler. Tvådagarsregeln minskar antalet falska alarm "
    "vid tillfälliga dippar."
)
brödtext(
    "Kategorierna 1–2 (Ankaret och Kassaflödet) säljs aldrig vid dipp — de köps mer av. "
    "Det är filosofin bakom strategin: förutsägbarhet minskar felaktiga beslut."
)


# ══════════════════════════════════════════════
#  3. OM AKTIER OCH PORTFÖLJFÖRVALTNING
# ══════════════════════════════════════════════
sidbyte()
rubrik1("3. Om aktier och portföljförvaltning")
ingress("Grunderna som appen bygger på")

brödtext(
    "Att investera i aktier är att äga en liten del av ett företag. Du delar i dess framgångar — "
    "och i dess misslyckanden. Det är inte ett spel, det är ett ägarskap. "
    "Och precis som ett ägarskap kräver det tålamod, disciplin och en plan."
)

rubrik2("Aktiemarknadens grundläggande natur")
brödtext(
    "Aktiemarknadens långsiktiga trend har historiskt gått uppåt — S&P 500 har levererat "
    "ungefär 10 % per år i genomsnitt sedan 1957, justerat för inflation ungefär 7 %. "
    "Men vägen dit har innehållit krascher på 30–50 % och perioder av tio år utan avkastning. "
    "Det är priset du betalar för den långsiktiga avkastningen."
)
citat("Time in the market beats timing the market.")
brödtext(
    "Det är den viktigaste meningen i hela investeringsvärlden. Och det är skälet till att "
    "Ankaret-kategorin — indexfonder — aldrig säljs vid dipp. De håller dig exponerad mot "
    "uppgången oavsett vad kortsiktigt brus säger."
)

rubrik2("Diversifiering — den enda gratislunchen")
brödtext(
    "Finansteori har bevisat att diversifiering minskar risk utan att offra avkastning, "
    "upp till en viss punkt. Att äga 20 aktier i olika branscher och valutor minskar "
    "bolagsspecifik risk dramatiskt. Men att äga 200 aktier ger marginellt bättre skydd — "
    "och enormt mycket mer komplexitet att hantera."
)
brödtext(
    "Strategiportföljens sex kategorier är designade för att ge meningsfull diversifiering "
    "utan att kräva ständig uppmärksamhet. Indexfonderna täcker tusentals bolag per automatik."
)

rubrik2("Psykologin — det svåraste verktyget")
brödtext(
    "Ingen investerares värsta fiende är marknaden. Det är investeraren själv. "
    "Rädsla säljer på botten. Girighet köper på toppen. Bekräftelsebias gör att vi letar "
    "efter information som stärker vad vi redan tror."
)
punkt("Rädsla och girighet är de vanligaste orsakerna till dåliga investeringsbeslut.")
punkt("En strategi med regler minskar utrymmet för impulsiva beslut.")
punkt("Beslutsloggen tvingar dig att motivera dina köp och sälj — vilket förbättrar beslutskvaliteten över tid.")

rubrik2("FX-motorn — bolagsvinst vs. valutavinst")
brödtext(
    "Ett unikt inslag i Strategiportföljen är FX-motorn: för utländska innehav separerar appen "
    "den vinst som kommer från bolagets faktiska kursutveckling från den vinst (eller förlust) "
    "som beror på att SEK försvagats eller stärkts mot USD, EUR eller NOK. "
    "Det ger en ärligare bild av vad ditt ägarskap faktiskt levererar."
)


# ══════════════════════════════════════════════
#  4. RISKER OCH VERKLIGHETEN
# ══════════════════════════════════════════════
sidbyte()
rubrik1("4. Risker och verkligheten")
ingress("Kapital på börsen är kapital du är beredd att förlora — det är inte en klyscha")

brödtext(
    "Det är möjligt att förlora hela det investerade kapitalet i enskilda aktier. "
    "Bolag kan gå i konkurs, sektorer kan kollapsa, makroekonomiska händelser kan radera "
    "decenniers uppgång på månader. Appen hjälper dig att se dessa risker — den eliminerar dem inte."
)

rubrik2("Typer av risk du tar")
punkt(
    "Enskilda bolag kan misslyckas. Enron, Lehman Brothers — de slog mot noll. "
    "Motmedel: diversifiering. Aldrig mer än 15 % i ett enskilt bolag.",
    "Bolagsrisk:"
)
punkt(
    "En hel bransch kan gå fel. Banker 2008, tech 2000. "
    "Motmedel: kategorierna täcker flera sektorer med olika riskprofil.",
    "Sektorrisk:"
)
punkt(
    "USD/SEK kan röra sig 15–20 % på ett år. FX-motorn visar exakt hur mycket av din "
    "avkastning som är valuta — inte bolag. MA200 mäts alltid i lokal valuta.",
    "Valutarisk:"
)
punkt(
    "Pengar i aktier är inte kontanter. Motmedel: investera aldrig kapital du behöver "
    "inom 3–5 år. Kassa-sektionen visar tillgänglig likviditet.",
    "Likviditetsrisk:"
)
punkt(
    "Du säljer i panik vid -30 % och köper tillbaka när marknaden är upp +40 %. "
    "Det är den vanligaste och dyraste misstaget. MA200-regeln ersätter känslan med en regel.",
    "Beteenderisk:"
)

rubrik2("Nödutgångar — när allt annat sviker")
brödtext(
    "Utöver MA200-regeln beräknar appen nödutgångar: om en aktie faller till 90 % av ditt "
    "genomsnittliga anskaffningsvärde (GAV), visas en röd nödutgångssignal. "
    "Det är en hård stopp för kategorierna 3–6, en mjuk analysvarning för 1–2."
)
brödtext(
    "Disciplin i nedsida är lika viktigt som disciplin i uppsida. Den som inte vet när de "
    "ska sälja, vet egentligen inte varför de köpt."
)
citat("The stock market is a device for transferring money from the impatient to the patient. — Warren Buffett")


# ══════════════════════════════════════════════
#  5. APPENS FUNKTIONER
# ══════════════════════════════════════════════
sidbyte()
rubrik1("5. Appens funktioner i korthet")
ingress("v3.13: Tio flikar, ett syfte — disciplinerad portföljförvaltning")

punkt(
    "Portföljöversikt med totalt värde, nyckeltal (nettoinsatt, avkastning, likviditet), "
    "FX-justerade signaler, kategorifördelning och värdeutvecklingsdiagram med periodväljare.",
    "▸  Dashboard"
)
punkt(
    "Alla positioner med MA200-avstånd i lokal valuta, bolagsvinst vs. valutavinst (FX-motor), "
    "tvådagarsstatus och nödutgångssignal. Redigerbar inline med MA200, ticker och kategori.",
    "▸  Innehav"
)
punkt(
    "Samlad vy med säljsignaler (röd/gul/grön), bevakningslista, nödutgångar, "
    "ombalanseringsbehov per kategori och hög koncentrationsrisk (>8 % av portfölj).",
    "▸  Signaler"
)
punkt(
    "Kategorifördelning med aktuellt värde, avkastning och portföljandel. "
    "Ombalanseringsassistent: 'Köp/Sälj för X kr' om vikt avviker >2 % från målintervall.",
    "▸  Kategorier"
)
punkt(
    "Komplett transaktionshistorik från Avanza. Utdelningar kopplas automatiskt till "
    "rätt aktie via ISIN. Filterbart per konto, typ och period.",
    "▸  Transaktioner"
)
punkt(
    "Förväntade utdelningar de kommande 90 dagarna baserat på historiska utdelningsmånader. "
    "Beräknas per aktie och summeras till total förväntad likviditetstillförsel.",
    "▸  Utdelningskalender"
)
punkt(
    "Veckovisa anteckningar om marknad och beslut. Exporterbar som del av Excel-backup.",
    "▸  Beslutslogg"
)
punkt(
    "Tillgängligt för köp per konto (ISK, KF, Sparkonto) hämtas automatiskt från "
    "positionsfilen. Manuell override möjlig. Sparkontot hanteras separat.",
    "▸  Kassa"
)
punkt(
    "Saldon per konto jämförda direkt mot Avanzas översikt: kontonummer, investerat värde, "
    "tillgängligt för köp och totalt per konto. Kontroll mot appens beräknade portföljvärde.",
    "▸  Avstämning mot Avanza"
)
punkt(
    "Kontokonfiguration, kategori-editor, strategiparametrar (MA200-gränser, nödutgång, "
    "ombalansering, konc.risk), profil, värdepappersfilter och export/import av inställningar som JSON.",
    "▸  Inställningar"
)
punkt(
    "Exporterar all portföljdata till Excel (7 ark: innehav, transaktioner, historik, "
    "kassa, kategorier, inställningar, register). Full återställning möjlig från backup.",
    "▸  Excel backup & återställning"
)

rubrik2("Import från Avanza")
brödtext(
    "Appen importerar tre CSV-filtyper från Avanza. Dra filerna direkt till rätt kort "
    "på Importera-fliken, eller välj flera filer samtidigt."
)
punkt("positioner_DATUM.csv — aktuella innehav med GAV, antal, marknadsvärde och kassa per konto")
punkt("transaktioner_DATUM.csv — köp, sälj, utdelningar och övriga transaktioner")
punkt("inkopskurs_DATUM.csv — engångsimport för att sätta historiska anskaffningsvärden")


# ══════════════════════════════════════════════
#  6. KOM IGÅNG PÅ FEM MINUTER
# ══════════════════════════════════════════════
sidbyte()
rubrik1("6. Kom igång på fem minuter")
ingress("Tre CSV-filer från Avanza — det är allt som krävs för att börja")

rubrik2("Steg 1: Öppna appen")
brödtext(
    "Gå till https://stockholmsvy1-droid.github.io/strategiportfoljen/ i Safari (iPad) "
    "eller valfri webbläsare på dator. Bokmärk sidan — appen kräver ingen installation."
)

rubrik2("Steg 2: Exportera från Avanza")
brödtext("Logga in på Avanza i webbläsaren (inte mobilappen). Gå till:")
punkt("Min ekonomi → Innehav → Exportera → positioner_DATUM.csv")
punkt("Min ekonomi → Transaktioner → Exportera transaktioner → transaktioner_DATUM.csv")
punkt("(Första gången) Min ekonomi → Innehav → Exportera inköpskurser → inkopskurs_DATUM.csv")

rubrik2("Steg 3: Importera filerna")
brödtext(
    "Gå till fliken Importera. Dra eller välj de CSV-filerna. "
    "Importera i ordning: transaktioner → positioner → inköpskurser. "
    "Appen läser dem och bygger upp innehav, transaktioner och kassinformation."
)

rubrik2("Steg 4: Sätt MA200-värden")
brödtext(
    "I Innehav-fliken — klicka på ett innehav och ange MA200 i lokal valuta. "
    "Hitta värdet via din mäklare eller finance.yahoo.com. "
    "Du kan också ange en ticker och hämta MA200 automatiskt via Alpha Vantage API."
)

rubrik2("Steg 5: Konfigurera konton (Inställningar)")
brödtext(
    "Gå till Inställningar → Kontokonfiguration och verifiera att dina Avanza-konton "
    "stämmer (namn, kontonummer, typ ISK/KF/SPAR). Kontrollera att sparkontots saldo "
    "är angivet i Kassa-fliken."
)

rubrik2("Steg 6: Stäm av mot Avanza")
brödtext(
    "Gå till Avstämning-fliken efter import. Summorna per konto ska stämma med "
    "Avanzas översikt. En differens under 1 % är normal (FX-timing). "
    "Större differens indikerar att något kan behöva kontrolleras."
)

rubrik2("Steg 7: Exportera backup varje vecka")
brödtext(
    "Gå till Importera → Backup & Återställning → Ladda ner Excel-backup. "
    "Spara filen i iCloud eller OneDrive. Det är din enda kopia av all data."
)


# ══════════════════════════════════════════════
#  7. NYHETER I VERSION 3.13
# ══════════════════════════════════════════════
sidbyte()
rubrik1("7. Nyheter i version 3.13")
ingress("Stabil grund — Avstämning mot Avanza fungerar nu fullt ut")

rubrik2("Avstämning omarbetad")
brödtext(
    "Fliken är ombyggd för att ge en tillförlitlig jämförelse mot Avanzas egna översiktsvy. "
    "Ny layout med framträdande summaryrad (Totalt värde + Tillgängligt för köp) i Avanzas stil. "
    "Varje kontorad visar nu investerat + kassa = exakt samma total som Avanza."
)
punkt("Avanza sparande Martin (SPAR) visas med korrekt saldo från manuellt angivet värde.")
punkt("Tillgängligt för köp per konto visas korrekt (visades felaktigt som 0 kr).")
punkt("Kontots totalt = investerat + kassa — matchar Avanzas kontovisning.")

rubrik2("Buggfix: Eget fondsparande (sedan v3.03)")
brödtext(
    "Kontot 'Eget fondsparande' var sedan version 3.03 felaktigt kopplat till ett pensionskonto "
    "(kontonr 9552-6014837) i stället för det riktiga ISK-kontot (9557-7346055). "
    "Felet medförde att Avanza Zero importerades med fel värde och att kontovärdena "
    "var ca 25 000 kr för låga i Avstämning."
)
punkt("Rätt kontonummer: 9557-7346055 (ISK, Eget fondsparande).")
punkt("Pensionskontot 9552-6014837 exkluderas nu korrekt från all import.")
punkt("Befintlig data migreras automatiskt vid laddning — ingen manuell åtgärd krävs.")

rubrik2("Arkitektur: manuell kassa ingår inte i portföljvärdet")
brödtext(
    "Manuella insättningar och uttag i Kassa-sektionen räknas inte längre in i portföljvärdet. "
    "De används enbart för att beräkna nettoinsatt kapital och avkastning. "
    "Ändringen eliminerar dubbelräkning: sparkontots saldo finns redan med i Avstämning "
    "— att också räkna in en insättning som registrerade att pengarna sattes in vore "
    "att räkna samma kapital två gånger."
)
punkt("Portföljvärde = innehav + Avanza-kassa (tillgängligt för köp från positionsfilen).")
punkt("Nettoinsatt-kortet på Dashboard visar fortfarande insättningshistoriken för avkastningsberäkning.")
punkt("Kontrolldifferensen i Avstämning mot app är nu nära noll (enbart FX/timing-skillnader).")

rubrik2("Tidigare versioner i korthet")
punkt("v3.12 — Kontohantering omskriven: kontonr som stabilt internt index eliminerar normNamn-ambiguitet.")
punkt("v3.11 — Ny Inställningar-sektion: kontokonfiguration, kategori-editor, strategiparametrar, profil, VP-filter.")
punkt("v3.10 — Interaktivt portföljutvecklingsdiagram, kontoregister, importordningsguide, förbättrad diff-wizard.")
punkt("v3.03 — Avstämningsflik med saldon per konto, kassa från positionsfilen, allow-list i import.")
punkt("v3.02 — Ny Avstämningsflik, periodbaserat nettoinsatt kapital.")

fotnot()

doc.save(dst)
print(f"Skapad: {dst}")
