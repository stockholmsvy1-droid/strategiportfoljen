"""
Skapar Strategiportfoljen_Beskrivning_v303.docx
Baserad på v302 — uppdaterad med v3.03-nyheter
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil, os

src = r"C:\Users\hejma\Projekt_Claude\strategiportfoljen\Strategiportfoljen_Beskrivning_v302.docx"
dst = r"C:\Users\hejma\Projekt_Claude\strategiportfoljen\Strategiportfoljen_Beskrivning_v303.docx"
dst_tmp = r"C:\Users\hejma\Projekt_Claude\strategiportfoljen\Strategiportfoljen_Beskrivning_v303_ny.docx"

NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0x0E, 0xA5, 0xE9)
MUTED  = RGBColor(0x6B, 0x72, 0x80)
BLACK  = RGBColor(0x11, 0x18, 0x27)

save_path = dst_tmp if os.path.exists(dst) else dst
shutil.copy2(src, save_path)
doc = Document(save_path)

# ── Ersätt versionsnummer i alla paragrafer ──
for para in doc.paragraphs:
    for run in para.runs:
        for old, new in [("v3.02","v3.03"),("v302","v303"),
                         ("version 3.02","version 3.03"),
                         ("Version 3.02","Version 3.03"),
                         ("3.02","3.03")]:
            if old in run.text:
                run.text = run.text.replace(old, new)

# ── Ersätt i tabeller ──
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if "3.02" in run.text:
                        run.text = run.text.replace("3.02","3.03")

# ── Hjälpfunktioner ──
def sidbyte(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(
        __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE
    )

def rubrik1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

def rubrik2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ACCENT

def brödtext(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.color.rgb = BLACK

def punkt(doc, text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    if prefix:
        r1 = p.add_run(prefix + " ")
        r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5); r.font.color.rgb = BLACK

# ── Ny sida med v3.03-nyheter ──
sidbyte(doc)
rubrik1(doc, "8. Nyheter i version 3.03")
rubrik2(doc, "Kontohantering, allow-list och buggfixar")
brödtext(doc,
    "Version 3.03 fokuserar på tillförlitlig kontoidentifiering och importkontroll — "
    "appen hanterar nu alla 6 Avanza-konton korrekt oavsett namnlikhet eller importordning."
)
rubrik2(doc, "Kontofiltrering — ny allow-list")
brödtext(doc,
    "Importlogiken är omskriven från block-list till allow-list. Enbart de 6 konfigurerade "
    "kontona importeras; övrig data ignoreras automatiskt utan manuell filtrering."
)
punkt(doc, "Inga okända konton kan glida in vid import, oavsett hur Avanza exporterar.")
punkt(doc, "Ny hjälptext förklarar allow-list-logiken direkt i appen.")

rubrik2(doc, "Sparkonto alltid synligt i Avstämning")
brödtext(doc,
    "Avanza Sparande Martin visas alltid i Saldon per konto-tabellen, "
    "även om inget saldo är angivet. Saknas saldo visas 'Ange saldo i Kassa-fliken'."
)

rubrik2(doc, "Kontonummer under kontonamn")
brödtext(doc,
    "Kontonumret visas under kontonamnet i Saldon per konto-tabellen — lättare "
    "att matcha mot rätt konto i Avanza."
)

rubrik2(doc, "Kassa-dropdown — alltid alla 6 konton")
brödtext(doc,
    "Väljarmenyn i Kassa-fliken visar alltid samtliga 6 konton, "
    "oavsett om en transaktionsfil är importerad."
)

rubrik2(doc, "Buggfix — exakt kontoidentifiering")
brödtext(doc,
    "Liknamnda konton (t.ex. '1. Utländska Aktier 2025' och '2. Utländska Aktier 2025') "
    "visades felaktigt som duplikat. Exakt matchning provas nu alltid först; "
    "normalisering används bara som reserv."
)

rubrik2(doc, "iOS-kompatibla importknappar")
brödtext(doc,
    "Importknapparna är bytta till label+for-element för pålitlig filväljare i Safari på iPad."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Strategiportföljen v3.03  ·  Byggt för Martin  ·  Strategi från januari 2026")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED

doc.save(save_path)
print(f"Skapad: {save_path}")
