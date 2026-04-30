"""
Skapar Admin_Forslag_Strategiportfoljen.docx
Analys och förslag på administrationsfunktioner för Strategiportföljen
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0x0E, 0xA5, 0xE9)
MUTED  = RGBColor(0x6B, 0x72, 0x80)
BLACK  = RGBColor(0x11, 0x18, 0x27)
GOLD   = RGBColor(0xCA, 0x8A, 0x04)
GREEN  = RGBColor(0x15, 0x80, 0x3D)
RED    = RGBColor(0xDC, 0x26, 0x26)

def sidbyte():
    p = doc.add_paragraph()
    p.add_run().add_break(
        __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE
    )

def linje():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E3A5F')
    pBdr.append(bottom)
    pPr.append(pBdr)

def rubrik1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

def rubrik2(text, color=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = color

def rubrik3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY

def bröd(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.color.rgb = BLACK

def punkt(text, bold_prefix=None, color=BLACK):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ":  ")
        r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5); r2.font.color.rgb = color

def tabell(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E3A5F')
        tcPr.append(shd)
    # Rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        fill = 'F1F5F9' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            r = cell.paragraphs[0].add_run(str(cell_text))
            r.font.size = Pt(9.5); r.font.color.rgb = BLACK
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)

def prioritet_badge(text, color_hex):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(f"  {text}  ")
    r.bold = True; r.font.size = Pt(9)
    rgb = tuple(int(color_hex[i:i+2], 16) for i in (0,2,4))
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ════════════════════════════════════════════════════════════════════
#  TITELSIDA
# ════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(50)
p.paragraph_format.space_after  = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Strategiportföljen")
r.bold = True; r.font.size = Pt(32); r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Administration & Inställningar")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
r = p.add_run("Analys och förslag — version 3.03 · April 2026")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = MUTED

linje()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
r = p.add_run(
    "Det här dokumentet analyserar vad som idag är hårdkodat i appen\n"
    "och föreslår konkreta administrationsfunktioner som gör Strategiportföljen\n"
    "mer generell och lätt att anpassa — utan att redigera kod."
)
r.font.size = Pt(11); r.font.color.rgb = MUTED

sidbyte()

# ════════════════════════════════════════════════════════════════════
#  1. INLEDNING
# ════════════════════════════════════════════════════════════════════

rubrik1("1. Inledning och syfte")
bröd(
    "Strategiportföljen är idag en välbyggd app som fungerar utmärkt för Martins specifika "
    "investeringsstrategi. Men ett antal viktiga parametrar — konton, kategorier, MA200-regler "
    "och tröskelvärden — är inbäddade direkt i koden. Det innebär att varje gång strategin "
    "förändras, eller om en annan person vill använda appen, krävs kodredigering."
)
bröd(
    "Målet med det här dokumentet är att identifiera exakt vad som behöver kunna konfigureras "
    "via ett användargränssnitt, och föreslå hur en dedikerad administrationsfunktion kan se ut. "
    "Appen ska alltid ha ett fungerande standardläge — men allt ska kunna justeras av användaren "
    "utan teknisk kunskap."
)

rubrik2("Vägledande principer")
punkt("Allt som ändras när strategin förändras ska kunna ändras i UI:t — aldrig i koden.",
      "Konfigurering utan kod")
punkt("Appen levereras med fabriksinställningar för Martins strategi. Ett klick återställer allt.",
      "Standardläge med återställning")
punkt("Inga data (innehav, historik, transaktioner) raderas när inställningar ändras.",
      "Inställningar påverkar inte data")
punkt("Administrationsgränssnittet ska vara begripligt på iPad utan teknisk förkunskap.",
      "iPad-first")

sidbyte()

# ════════════════════════════════════════════════════════════════════
#  2. NULÄGE — VAD ÄR HÅRDKODAT
# ════════════════════════════════════════════════════════════════════

rubrik1("2. Nuläge — vad är hårdkodat idag")
bröd(
    "Följande tabell visar alla element som idag kräver kodredigering för att ändras. "
    "Kolumnen 'Ändring' anger hur ofta detta kan behöva justeras i praktiken."
)

tabell(
    ["Element", "Nuvarande värde", "Typ", "Ändring"],
    [
        ["6 Avanza-konton", "Namn, kontonr, typ (ISK/KF/SPAR)", "Hårdkodad lista", "Vid byte av konto"],
        ["Kategorier (6 st)", "Namn, emoji, färg, mål, signal", "Delvis konfig.", "Vid strategibyte"],
        ["Kategori-CRUD", "Finns men via prompt()-dialoger", "Primitivt UI", "Ofta"],
        ["MA200 grön-gräns", ">5% ovanför MA200", "Hårdkodat tal", "Sällan"],
        ["MA200 gul-gräns", "±5% kring MA200", "Hårdkodat tal", "Sällan"],
        ["Gummiband gul-zon", "15–40% ovanför MA200", "Hårdkodat tal", "Sällan"],
        ["Gummiband orange-zon", ">40% ovanför MA200", "Hårdkodat tal", "Sällan"],
        ["Nödutgång hård", "Kurs < 90% av GAV (kat 3–6)", "Hårdkodat tal", "Sällan"],
        ["Nödutgång mjuk", "Kurs < 90% av GAV (kat 1–2)", "Hårdkodat tal", "Sällan"],
        ["Ombalanseringströskel", ">2 procentenheters avvikelse", "Hårdkodat tal", "Sällan"],
        ["Tvådagarsregel", "På för kat 3–6", "Hårdkodad logik", "Sällan"],
        ["Exkl. värdepapper", "['zomedica']", "Hårdkodad lista", "Ibland"],
        ["Exkl. konton", "['9557-7346055'] (pension)", "Hårdkodad lista", "Sällan"],
        ["Strateginamn", "'Aktiestrategi 2026'", "Hårdkodad text", "Varje år"],
        ["Portfolioägarens namn", "'Martin'", "Hårdkodad text", "Vid överlåtelse"],
    ],
    col_widths=[4.2, 4.5, 3.2, 2.6]
)

bröd(
    "Kategori-objektet (`KATEGORIER`) har redan en bra struktur med `DEFAULT_KATEGORIER` "
    "och `sparaKategorier()` / `laddaKategorier()` — grundarkitekturen för konfigurering finns. "
    "Det som saknas är ett ordentligt användargränssnitt och att utöka mönstret till övriga parametrar."
)

sidbyte()

# ════════════════════════════════════════════════════════════════════
#  3. FÖRSLAG A — KATEGORI-EDITOR
# ════════════════════════════════════════════════════════════════════

rubrik1("3. Förslag A — Kategori-editor")
bröd(
    "Den befintliga kategori-CRUD:en (skapaKategori, redigeraKategori, taBortKategori) "
    "använder webbläsarens inbyggda prompt()-dialoger. Det är primitivt, svårt att använda "
    "på iPad och ger dålig visuell återkoppling. Förslaget är ett dedikerat, visuellt "
    "redigeringsgränssnitt."
)

rubrik2("Gränssnittsbeskrivning")
bröd(
    "Kategori-editorn visas som ett kort per kategori, ordnade i en vertikal lista. "
    "Varje kort visar kategorins nuvarande utseende och har ett inbyggt redigeringsformulär "
    "som fälls ut vid klick på 'Redigera'."
)

rubrik3("Per kategori-kort:")
punkt("Namn (fritext, max 20 tecken)", "Namn")
punkt("Undertitel / typ-beskrivning (fritext, max 20 tecken)", "Undertitel")
punkt("Emoji — välj från en fördefinierad lista med ca 30 relevanta emojis", "Emoji")
punkt("Färg — color picker med 12 fördefinierade palettalternativ + fri hex-inmatning", "Färg")
punkt("Målvikt min % och max % — numerisk inmatning eller slider (0–100)", "Målvikt")
punkt("Signaltyp — radioknapp: Ingen signal | MA200 med tvådagarsregel", "Signal")
punkt("Beskrivning — fritext (visas i hjälptext och strategi-export)", "Beskrivning")

rubrik3("Listans knappar:")
punkt("Flytta upp / flytta ner — byt ordning (visas i Innehav, Dashboard, Signaler)", "Ordning")
punkt("Ta bort — kräver bekräftelse. Innehav i borttagen kategori omtilldelas kat 2.", "Ta bort")
punkt("Lägg till kategori — tomt kort läggs till sist i listan", "Lägg till")
punkt("Återställ alla kategorier — återgår till fabriksinställningarna för Martins strategi", "Återställ")

rubrik2("Konsekvenser av kategoriändringar")
bröd(
    "Eftersom kategorier refereras från varje innehav (h.kategori = 1–6) är det viktigt "
    "att hantera edge-cases:"
)
punkt("Om en kategori tas bort: alla innehav med den kategorin tilldelas automatiskt kategori 2 (Kassaflödet / default).")
punkt("Om ordningen ändras: kategori-ID:n (1–6) förblir oförändrade — det är bara visningsordningen som ändras.")
punkt("Om målvikterna inte summerar till 100%: en varning visas men spara tillåts — appen hanterar vikter relativt.")

rubrik2("Validering")
tabell(
    ["Fält", "Validering", "Felmeddelande"],
    [
        ["Namn", "1–20 tecken, unik bland aktiva kategorier", "Namn krävs och måste vara unikt"],
        ["Målvikt min", "0–100, ≤ målvikt max", "Min måste vara ≤ max"],
        ["Målvikt max", "0–100, ≥ målvikt min", "Max måste vara ≥ min"],
        ["Summa min", "Varning om summa min > 100%", "Varning: vikter överstiger 100%"],
        ["Antal kategorier", "Minst 1, max 12", "Minst en kategori krävs"],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

sidbyte()

# ════════════════════════════════════════════════════════════════════
#  4. FÖRSLAG B — STRATEGIPARAMETRAR
# ════════════════════════════════════════════════════════════════════

rubrik1("4. Förslag B — Strategiparametrar")
bröd(
    "En rad numeriska trösklar styr appens signal- och rebalanseringslogik. Alla är idag "
    "hårdkodade i JavaScript. Förslaget är att exponera dem i ett enkelt formulär med "
    "förklarande text och möjlighet att återgå till defaultvärdet per parameter."
)

rubrik2("Konfigurerbara parametrar")
tabell(
    ["Parameter", "Default", "Enhet", "Förklaring"],
    [
        ["Ombalanseringströskel", "2", "%", "Avvikelse från målvikt som triggar ombalanseringssignal på Dashboard"],
        ["MA200 — grön-gräns", "5", "% ovan MA200", "Kurs mer än X% över MA200 → grön signal (håll)"],
        ["MA200 — gul-gräns", "5", "% under MA200", "Kurs upp till X% under MA200 → gul (dag 1, bevaka)"],
        ["Gummiband — gul-zon", "15", "% ovan MA200", "Kurs X–Y% över MA200 → sträckt (gul gummiband)"],
        ["Gummiband — orange-zon", "40", "% ovan MA200", "Kurs >X% över MA200 → kraftigt sträckt (orange)"],
        ["Nödutgång hård (kat 3–6)", "10", "% under GAV", "Säljsignal om kurs faller mer än X% under GAV"],
        ["Nödutgång mjuk (kat 1–2)", "10", "% under GAV", "Varningssignal (ej säljsignal) vid samma gräns"],
        ["Tvådagarsregel", "På", "På/Av", "Kräver 2 stängningar under MA200 för röd signal"],
    ],
    col_widths=[4.5, 2.0, 3.0, 5.0]
)

rubrik2("Gränssnittsbeskrivning")
bröd(
    "Varje parameter visas som en rad med: fältnamn, numerisk input (eller toggle), "
    "enhetsetikett, och en liten '?' som visar förklaringstexten vid hover/klick. "
    "Till höger om varje fält: ett litet 'Återställ'-kryss som återgår till defaultvärdet "
    "för just den parametern. Fält med ändrat värde markeras med en blå kant."
)

rubrik2("Konsekvenser")
punkt(
    "Ändring av MA200-gränser påverkar omedelbart signalstatus för alla innehav — "
    "ingen re-import krävs.",
    "Direkt effekt"
)
punkt(
    "Ändring av ombalanseringströskel påverkar hur många kategorier som flaggas på Dashboard.",
    "Dashboard"
)
punkt(
    "Tvådagarsregel Av innebär att en enstaka stängning under MA200 räcker för röd signal.",
    "Tvådagarsregel"
)

sidbyte()

# ════════════════════════════════════════════════════════════════════
#  5. FÖRSLAG C — KONTOKONFIGURATION
# ════════════════════════════════════════════════════════════════════

rubrik1("5. Förslag C — Kontokonfiguration")
bröd(
    "Idag är Avanza-kontonamnen och kontonumren hårdkodade som en JavaScript-konstant "
    "(AVANZA_KONTON_LIST). Det gör att om Martin öppnar ett nytt konto, byter namn på "
    "ett konto, eller om en annan person vill använda appen — krävs kodredigering. "
    "Det är den förändring som har störst påverkan på appens generalitet."
)

rubrik2("Nuvarande konfiguration (fabriksinställning)")
tabell(
    ["Kontonamn", "Kontonummer", "Typ", "Ordning"],
    [
        ["1. Sverige nov 2025", "1914035", "ISK", "1"],
        ["2. Norden nov 2025", "1080302", "ISK", "2"],
        ["Eget fondsparande", "6014837", "ISK", "3"],
        ["1. Utländska Aktier 2025", "7882604", "KF", "4"],
        ["2. Utländska Aktier 2025", "1913799", "KF", "5"],
        ["Avanza sparande Martin", "40080455", "SPAR", "6"],
    ],
    col_widths=[5.0, 3.0, 2.0, 2.0]
)

rubrik2("Gränssnittsbeskrivning")
punkt("Lista alla konfigurerade konton med namn, kontonummer och typ", "Visa")
punkt("Fritext för namn, numerisk inmatning för kontonummer", "Lägg till konto")
punkt("Ändra namn eller kontonummer direkt i listan", "Redigera")
punkt("Kräver bekräftelse. Tar inte bort data — bara kontokopplingen.", "Ta bort konto")
punkt("Pil-knappar styr i vilken ordning konton visas i Kassa och Avstämning", "Ordning")
punkt("Typ: ISK / KF / SPAR — styr om kontot visas i Kassa-tabellen eller bara Avstämning", "Typ")

rubrik2("Viktiga konsekvenser")
bröd(
    "Kontokonfigurationen styr vilka konton som accepteras vid CSV-import (allow-list). "
    "Att ändra kontonummer påverkar därför direkt vad som importeras nästa gång. "
    "Befintliga positionsvärden i localStorage påverkas inte retroaktivt."
)
punkt(
    "Om ett konto läggs till: det börjar importeras nästa gång positionsfilen importeras.",
    "Tillägg"
)
punkt(
    "Om ett konto tas bort: importerade data för det kontot finns kvar i historiken, "
    "men kontot importeras inte längre.",
    "Borttagning"
)
punkt(
    "Exkluderade konton (pension m.fl.) hanteras separat i Filter-sektionen.",
    "Exkludering"
)

sidbyte()

# ════════════════════════════════════════════════════════════════════
#  6. FÖRSLAG D — VÄRDEPAPPERSFILTER
# ════════════════════════════════════════════════════════════════════

rubrik1("6. Förslag D — Värdepappersfilter")
bröd(
    "Idag filtreras Zomedica alltid bort (hårdkodat i EXKLUDERA_VP). "
    "Pensionskontot är exkluderat via EXKLUDERA_KONTON. "
    "Förslaget är att göra dessa listor redigerbara via UI."
)

rubrik2("Exkluderade värdepapper")
bröd(
    "En lista med värdepapper som alltid ignoreras vid CSV-import, oavsett konto. "
    "Används för aktier man äger men inte vill följa i strategiappen (pensionssparande, "
    "enskilda innehav i exkluderade konton, etc.)."
)
punkt("Fritext för värdepappersnamn (matchar case-insensitivt mot CSV:ets Namn-kolumn)", "Lägg till")
punkt("Ta bort ur listan med ett klick", "Ta bort")
punkt("Valfri kommentar: 'Varför exkluderas detta?'", "Anledning")

rubrik2("Exkluderade konton")
bröd(
    "En lista med kontonummer som alltid ignoreras vid import, oavsett om de råkar "
    "matcha ett namn i kontolistan. Primärt för pensionskonton."
)
punkt("Numerisk inmatning av kontonummer", "Lägg till")
punkt("Valfri etikett: 'Pensionskonto', 'Utländskt konto' etc.", "Etikett")

sidbyte()

# ════════════════════════════════════════════════════════════════════
#  7. FÖRSLAG E — STRATEGI-INFO & PROFIL
# ════════════════════════════════════════════════════════════════════

rubrik1("7. Förslag E — Strategi-information och profil")
bröd(
    "Appen visar idag 'Aktiestrategi 2026' och 'Martin' som hårdkodade texter på "
    "flera ställen — i hjälptexter, i export-HTML och i dokumentmallar. "
    "Förslaget är att göra dessa till redigerbara fält."
)

rubrik2("Redigerbara fält")
tabell(
    ["Fält", "Default", "Används i"],
    [
        ["Portföljnamn", "Strategiportföljen", "Sidhuvud, export, hjälptexter"],
        ["Ägarens namn", "Martin", "Dashboard, export-HTML, dokumentmallar"],
        ["Strateginamn", "Aktiestrategi 2026", "Strategi-hjälptext, export"],
        ["Strategibeskrivning", "Regelbaserad förvaltning...", "Hjälptext, export"],
        ["Strategi startdatum", "Januari 2026", "Hjälptext, historik"],
        ["Senast uppdaterad", "5 april 2026", "Strategi-hjälptext"],
    ],
    col_widths=[4.0, 4.0, 6.5]
)

rubrik2("Övriga profilinställningar")
punkt("Ljust / Mörkt tema — finns redan, men kan samlas här", "Tema")
punkt("Primärvaluta (just nu alltid SEK) — för framtida internationalitet", "Valuta")
punkt("Alpha Vantage API-nyckel — finns redan, kan samlas här", "API-nyckel")

sidbyte()

# ════════════════════════════════════════════════════════════════════
#  8. FÖRSLAG F — INSTÄLLNINGAR-SEKTION I APPEN
# ════════════════════════════════════════════════════════════════════

rubrik1("8. Sammanslagning — Inställningar-sektionen")
bröd(
    "Alla ovanstående förslag samlas i en ny sektion '⚙️ Inställningar' i appens "
    "navigering. Sektionen delas in i flikar. En tydlig 'Återställ fabriksinställningar'-knapp "
    "finns alltid tillgänglig och återställer ALLT till Martins ursprungsstrategi utan att "
    "röra portföljdata (innehav, historik, transaktioner)."
)

rubrik2("Navigation och struktur")
tabell(
    ["Flik", "Innehåller", "Prioritet"],
    [
        ["Kategorier", "Visuell kategori-editor (Förslag A)", "★★★ Hög"],
        ["Strategi", "Parametrar (Förslag B) + Strategiinfo (Förslag E)", "★★ Medium"],
        ["Konton", "Kontokonfiguration (Förslag C)", "★★★ Hög"],
        ["Filter", "VP-filter + kontouteslutning (Förslag D)", "★ Låg"],
        ["Data", "Datarensning + JSON-export av inställningar", "★ Låg"],
    ],
    col_widths=[3.0, 8.5, 3.0]
)

rubrik2("Standardläge och återställning")
bröd(
    "Appen levereras med Martins strategi som fabriksinställning. Begreppet 'standardläge' "
    "innebär att:"
)
punkt(
    "Alla inställningsobjekt har ett DEFAULT-värde inbakat i koden — "
    "detta är appens facit och kan aldrig förstöras.",
    "Oförstörbart default"
)
punkt(
    "Fält som avviker från default markeras visuellt (blå kant, liten 'ändrad'-etikett) "
    "så att användaren alltid vet vad som är anpassat.",
    "Tydlig indikation"
)
punkt(
    "Per fält: liten 'Återställ'-knapp återgår till defaultvärdet för just det fältet.",
    "Granulär återställning"
)
punkt(
    "Global 'Återställ alla inställningar'-knapp med bekräftelsedialog återställer "
    "kategorier, parametrar, konton och filter — men ALDRIG innehav, historik eller transaktioner.",
    "Global återställning"
)

rubrik2("Export och import av inställningar")
bröd(
    "Inställningarna (kategorier, parametrar, konton, filter) kan exporteras som en "
    "JSON-fil och importeras på en annan enhet. Det gör det möjligt att:"
)
punkt("Flytta sin anpassade strategi till en ny enhet utan att göra om allt")
punkt("Dela sin strategikonfiguration med en annan Avanza-användare")
punkt("Ha ett 'säkerhetskopia av strategin' separat från portföljdatan")

sidbyte()

# ════════════════════════════════════════════════════════════════════
#  9. PRIORITERING OCH GENOMFÖRANDEORDNING
# ════════════════════════════════════════════════════════════════════

rubrik1("9. Prioritering och genomförandeordning")
bröd(
    "Nedanstående ordning baseras på värde för användaren vägt mot implementationskomplexitet. "
    "Varje steg är oberoende och kan implementeras och testas separat."
)

rubrik2("Steg 1 — Kontokonfiguration  (1–2 dagars arbete)")
bröd(
    "Avgörande för generalitet. Utan detta kräver varje ny Avanza-användare kodredigering. "
    "Flytta AVANZA_KONTON_LIST från hårdkodad konstant till data.kontokonfiguration i "
    "localStorage. Bygg enkel lista-UI med add/remove/edit."
)

rubrik2("Steg 2 — Kategori-editor  (2–3 dagars arbete)")
bröd(
    "Grundstrukturen finns (DEFAULT_KATEGORIER, sparaKategorier). Det som saknas är "
    "ett visuellt formulär istället för prompt()-dialoger. Relativt låg teknisk risk "
    "eftersom datamodellen redan är korrekt."
)

rubrik2("Steg 3 — Strategiparametrar  (1 dag)")
bröd(
    "Skapa ett DEFAULT_STRATEGI_PARAMS-objekt med alla trösklar. Ersätt de hårdkodade "
    "talen i ma200Signal(), gummibandZon() och renderSignaler() med referenser till "
    "objektet. Bygg ett enkelt formulär med sliders och reset-knappar."
)

rubrik2("Steg 4 — Strategi-info & Profil  (0.5 dag)")
bröd(
    "Relativt enkelt — spara ett objekt med namn, titel etc. i localStorage och ersätt "
    "hårdkodade strängar med references till objektet."
)

rubrik2("Steg 5 — Värdepappersfilter  (0.5 dag)")
bröd(
    "Flytta EXKLUDERA_VP och EXKLUDERA_KONTON till localStorage. Bygg enkel lista-UI."
)

rubrik2("Steg 6 — Export/import av inställningar  (0.5 dag)")
bröd(
    "Samla alla inställningsobjekt i ett JSON-paket. Knapp för download och file-input "
    "för import. Enkel att implementera sist när övriga steg är klara."
)

tabell(
    ["Steg", "Funktion", "Arbete", "Värde", "Risk"],
    [
        ["1", "Kontokonfiguration", "1–2 dagar", "★★★ Hög", "Låg"],
        ["2", "Kategori-editor", "2–3 dagar", "★★★ Hög", "Låg"],
        ["3", "Strategiparametrar", "1 dag", "★★ Medium", "Låg"],
        ["4", "Strategi-info & Profil", "0.5 dag", "★ Låg", "Minimal"],
        ["5", "Värdepappersfilter", "0.5 dag", "★ Låg", "Minimal"],
        ["6", "Export/import inställningar", "0.5 dag", "★ Låg", "Minimal"],
    ],
    col_widths=[1.5, 4.5, 2.5, 2.5, 2.0]
)

bröd(
    "Total uppskattad implementationstid: 5–8 dagars arbete för full implementation. "
    "Steg 1–2 ger störst utväxling och rekommenderas som första leverans."
)

# Sidfot
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Strategiportföljen  ·  Admin & Inställningar — Förslag  ·  April 2026  ·  Byggt för Martin")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED

# Spara
import os
dst = r"C:\Users\hejma\Projekt_Claude\strategiportfoljen\Admin_Forslag_Strategiportfoljen.docx"
doc.save(dst)
print(f"Skapad: {dst}")
