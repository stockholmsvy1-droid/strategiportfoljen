"""
Skapar Strategiportfoljen_Beskrivning_v208.docx
5-sidig beskrivning av appen, investeringsfilosofi och risker
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Sidmarginaler ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Färger ─────────────────────────────────────────────────
NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0x0E, 0xA5, 0xE9)
MUTED  = RGBColor(0x6B, 0x72, 0x80)
BLACK  = RGBColor(0x11, 0x18, 0x27)
GOLD   = RGBColor(0xCA, 0x8A, 0x04)

# ── Hjälpfunktioner ────────────────────────────────────────
def rubrik1(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = color
    return p

def rubrik2(text, color=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    return p

def rubrik3(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = color
    return p

def brödtext(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    return p

def kursiv(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.italic    = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = MUTED
    return p

def punkt(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = NAVY
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = BLACK
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLACK
    return p

def linje():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E3A5F')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def sidbyte():
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(__import__('docx.enum.text', fromlist=['WD_BREAK']).WD_BREAK.PAGE)

def citat(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.2)
    p.paragraph_format.right_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(f'"{text}"')
    run.italic    = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    return p

def not_text(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.7)
    run = p.add_run("⚠  " + text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = GOLD
    return p


# ════════════════════════════════════════════════════════════
#  TITELSIDA
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after  = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Strategiportföljen")
run.bold      = True
run.font.size = Pt(28)
run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("version 2.08")
run.font.size = Pt(14)
run.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
run = p.add_run("The Magnificent Martin's Money-Making Machine")
run.italic    = True
run.font.size = Pt(12)
run.font.color.rgb = MUTED

linje()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
run = p.add_run(
    "En personlig portföljapp för att följa en aktieportfölj med fokus\n"
    "på långsiktig tillväxt, riskmedvetenhet och strategisk disciplin.\n\n"
    "Byggt för iPad och dator · Öppnas direkt i webbläsaren · Inga konton"
)
run.font.size = Pt(11)
run.font.color.rgb = MUTED

sidbyte()


# ════════════════════════════════════════════════════════════
#  SIDA 1 — VAD ÄR STRATEGIPORTFÖLJEN?
# ════════════════════════════════════════════════════════════
rubrik1("1. Vad är Strategiportföljen?")
rubrik2("v2.08: Version med åtta rörliga delar — precis som en bra portfölj")

brödtext(
    "Strategiportföljen är en personlig webbapp byggd för att ge en samlad, "
    "överskådlig bild av en aktieportfölj som följer en genomtänkt investeringsstrategi. "
    "Appen ersätter kalkylblad och anteckningshäften med ett levande instrument — "
    "ett cockpit för den som vill förstå vad pengarna gör, inte bara titta på siffror."
)
brödtext(
    "Tanken bakom appen är enkel: du ska kunna öppna den på måndag morgon, "
    "importera tre CSV-filer från Avanza, och direkt se om något behöver din uppmärksamhet. "
    "Är du under MA200 på NVIDIA? Har SAAB gett signal två dagar i rad? "
    "Vad är den faktiska avkastningen i dollar jämfört med SEK? "
    "Det är de frågorna appen är byggd för att svara på — snabbt, korrekt och utan brus."
)

rubrik2("Arkitektur och filosofi")
brödtext(
    "Appen är avsiktligt byggt som en enda HTML-fil — index.html — utan backend, "
    "utan databas, utan inloggning. All data sparas i webbläsarens lokala lagring (localStorage). "
    "Det innebär att data aldrig lämnar din enhet, att appen fungerar offline, "
    "och att det inte finns någon server som kan drabbas av driftstopp eller dataintrång."
)
brödtext(
    "Det är ett medvetet val som speglar investeringsfilosofin: enkelhet minskar fel. "
    "Precis som en portfölj med för många innehav bli svår att övervaka, "
    "blir en app med för många beroenden svår att underhålla och lita på."
)

rubrik3("Appen används primärt på tre sätt:")
punkt("Veckovis uppföljning", "→")
punkt("(Måndag efter stängning) Importera positioner, transaktioner och eventuell kassa från Avanza. Granska nyckeltal och MA200-signaler.")
punkt("Beslutsunderlag", "→")
punkt("Innan ett köp eller sälj — verifiera kategoritillhörighet, MA200-avstånd och portföljvikt.")
punkt("Strategisk reflektion", "→")
punkt("Beslutsloggen och värdeutvecklingsdiagrammet ger en historisk bild av vad som fungerat och inte.")

sidbyte()


# ════════════════════════════════════════════════════════════
#  SIDA 2 — INVESTERINGSSTRATEGIN
# ════════════════════════════════════════════════════════════
rubrik1("2. Investeringsstrategin bakom appen")
rubrik2("v2.08: Åtta testsviter, sex kategorier — strukturen är allt")

brödtext(
    "Appen är designad kring en specifik investeringsstrategi med sex kategorier. "
    "Varje kategori representerar en roll i portföljen — en funktion, inte bara ett tema. "
    "Grundtanken är att en portfölj är som ett lag: du behöver en försvarsspelare, "
    "en målskytt och en som springer inkasso. Ingen av dem är bättre än den andra, "
    "men alla behövs."
)

rubrik3("De sex kategorierna")

kategorier = [
    ("⚓ Ankaret", "Indexfonder", "35–40 %",
     "Basen i portföljen. Köps regelbundet, säljs aldrig vid dipp. "
     "Ger exponering mot bred marknadsuppgång med minimal risk och låg kostnad. "
     "Utan ett stadigt ankare kan hela portföljen driva iväg i storm."),
    ("💰 Kassaflödet", "Utdelningsaktier", "20–25 %",
     "Stabila bolag med lång utdelningshistorik. Utdelningarna återinvesteras. "
     "Säljs inte vid marknadskorrigeringar — de är portföljens \u201alunga\u2019, "
     "som andas lugnt när resten av marknaden hyperventilerar."),
    ("🔵 Infrastrukturen", "AI-hårdvara / chip", "15–20 %",
     "NVIDIA, TSMC, AMD — bolag som bygger fundamentet för nästa teknologivåg. "
     "Här gäller MA200-regeln med tvådagarsbekräftelse. Volatilt men strukturellt starkt."),
    ("🧠 Hjärnan", "AI-mjukvara", "10–15 %",
     "Bolag som tjänar pengar på att AI faktiskt används — Microsoft, Palantir, Salesforce. "
     "Högre värdering men exponering mot en sekulär tillväxttrend."),
    ("🛡 Skölden", "Försvarsindustri", "5–10 %",
     "SAAB, Rheinmetall, BAE Systems. Defensiv exponering mot ett strukturellt "
     "underinvesterat segment. MA200-reglerna gäller."),
    ("✨ Berättelserna", "Kryddor / teman", "0–5 %",
     "Spekulativa positioner med hög conviction. Liten andel, stor potentiell uppsida "
     "— och full risk för förlust. Behandlas som lotter: man kan förlora allt, "
     "men man väljer medvetet hur mycket man är beredd att satsa."),
]

for namn, undertitel, vikt, beskrivning in kategorier:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"{namn}  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(f"({undertitel} · målvikt {vikt})")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = MUTED

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.2)
    p2.paragraph_format.space_after = Pt(4)
    r3 = p2.add_run(beskrivning)
    r3.font.size = Pt(10)
    r3.font.color.rgb = BLACK

rubrik2("MA200-regeln — din disciplinerade exit")
brödtext(
    "För kategorierna 3–6 gäller en tydlig regel: om en aktie stänger under sitt "
    "200-dagars glidande medelvärde (MA200) i lokal valuta, och det bekräftas "
    "ett andra handelsdag i rad — är det en säljsignal. Inte en panik-signal. "
    "En säljsignal. Du bestämmer fortfarande, men appen flaggar det."
)
brödtext(
    "MA200 anges alltid i lokal valuta — NVIDIA i USD, SAAB i SEK. "
    "Det eliminerar bruset från SEK-fluktuationer som annars ger falska signaler "
    "när kronan rör sig mot dollarn."
)

sidbyte()


# ════════════════════════════════════════════════════════════
#  SIDA 3 — OM AKTIEHANDEL GENERELLT
# ════════════════════════════════════════════════════════════
rubrik1("3. Om aktier och portföljförvaltning")
rubrik2("v2.08: Åtta ark i backupen, tusen sätt att förlora pengar — men fler att tjäna dem")

brödtext(
    "Att investera i aktier är att äga en liten del av ett företag. "
    "Du delar i dess framgångar — och i dess misslyckanden. "
    "Det är inte gambling, men det är inte heller en sparränta. "
    "Det är en avtalsrelation med en okänd framtid, och det är viktigt att förstå det innan man börjar."
)

rubrik3("Aktiemarknadens grundläggande natur")
brödtext(
    "Aktiemarknadens långsiktiga trend har historiskt gått uppåt — "
    "S&P 500 har levererat ungefär 10 % per år i genomsnitt sedan 1957, "
    "justerat för inflation ca 7 %. Men den trenden är inte jämn. "
    "Det finns år med -40 %, kvartal med panik och månader av brutal korrigering. "
    "Den som förblir investerad och inte säljer i panik har historiskt belönats. "
    "Den som försöker tajma marknaden har historiskt misslyckats — "
    "även professionella fondförvaltare slår sällan index konsekvent."
)

citat("Time in the market beats timing the market.")

brödtext(
    "Det är den viktigaste meningen i hela investeringsvärlden. "
    "Och det är skälet till att Ankaret-kategorin — indexfonder — "
    "aldrig säljs vid en dipp. Du väntar ut stormen. Du låter tid göra jobbet."
)

rubrik3("Diversifiering — den enda gratislunchen")
brödtext(
    "Finansteori har bevisat att diversifiering minskar risk utan att offra avkastning, "
    "upp till en viss punkt. Att äga 20 aktier i olika branscher ger nästan lika låg "
    "volatilitet som att äga 500 — men att äga 2 ger dramatiskt mycket högre risk än att äga 20. "
    "Strategiportföljens sex kategorier syftar till att sprida risk geografiskt, "
    "sektorsmässigt och tidsmässigt — kortsiktig tillväxt, långsiktig stabilitet, "
    "utdelningar idag och potential imorgon."
)

rubrik3("Psykologin — det svåraste verktyget")
brödtext(
    "Ingen investerares värsta fiende är marknaden. Det är investeraren själv. "
    "Rädsla säljer på botten. Girighet köper på toppen. "
    "Confirmation bias leder till att man bara läser det man redan tror. "
    "Appen motverkar det genom att vara objektiv: MA200 är vad den är, "
    "portföljvikten är vad den är, avkastningen är vad den är. "
    "Det finns ingen känsla i ett diagram — bara data."
)

punkt("Rädsla och girighet är de vanligaste orsakerna till dåliga investeringsbeslut.")
punkt("En strategi med regler minskar utrymmet för impulsiva beslut.")
punkt("Beslutsloggen tvingar dig att motivera dina köp och sälj — vilket förbättrar beslutskvaliteten över tid.")

rubrik3("Utdelningar — pengarna som arbetar medan du sover")
brödtext(
    "Utdelningsaktier delar ut en del av vinsten direkt till aktieägarna, "
    "vanligen en till fyra gånger per år. En aktie som ger 4 % direktavkastning "
    "och återinvesteras varje år fördubblar din position på 18 år via ränta-på-ränta — "
    "utan att aktiekursen rört sig ett öre. "
    "Kassaflödes-kategorin är byggd för att dra nytta av just detta."
)

sidbyte()


# ════════════════════════════════════════════════════════════
#  SIDA 4 — RISKER OCH VERKLIGHETEN
# ════════════════════════════════════════════════════════════
rubrik1("4. Risker och verkligheten")
rubrik2("v2.08: Åtta testfall i T4 — för att bekräfta att data inte försvinner. Kapital kan det.")

not_text(
    "Pengar du placerar på börsen är pengar du är beredd att förlora. "
    "Det är inte en klyscha — det är ett kontrakt med verkligheten."
)

brödtext(
    "Det är möjligt att förlora hela det investerade kapitalet i enskilda aktier. "
    "Bolag kan gå i konkurs, sektorer kan kollapsa, makroekonomiska kriser kan "
    "halvera en portfölj på månader. Ingen strategi — inte ens den bästa — "
    "eliminerar marknadsrisk. Strategin minimerar den."
)

rubrik3("Typer av risk du tar")
risker = [
    ("Bolagsrisk",
     "Enskilda bolag kan misslyckas. Enron, Lehman Brothers, Nokia — de slog mot noll. "
     "Motmedel: diversifiering. Aldrig mer än 5–10 % av portföljvärdet i ett enda innehav."),
    ("Sektorrisk",
     "En hel bransch kan gå fel. Banker 2008, tech 2000, energi 2015. "
     "Motmedel: kategorierna täcker flera sektorer med olika konjunkturkänslighet."),
    ("Valutarisk",
     "USD/SEK kan röra sig 15–20 % på ett år. En NVIDIA-position kan se ut att ha gått upp "
     "i SEK medan den gått ned i USD. "
     "Motmedel: FX-motorn i appen separerar bolagsvinst från valutavinst."),
    ("Likviditetsrisk",
     "Pengar i aktier är inte kontanter. Du kan inte ta ut dem på en sekund. "
     "Motmedel: investera aldrig kapital du behöver inom 1–3 år."),
    ("Ränterisk",
     "Höga räntor pressar aktievärderingar, särskilt i tillväxtsektorer. "
     "Motmedel: en del av portföljen i utdelningsaktier och realtillgångar."),
    ("Beteenderisk",
     "Du säljer i panik vid -30 % och köper tillbaka när marknaden är upp +40 %. "
     "Det är den vanligaste och dyraste misstaget. "
     "Motmedel: en strategi med regler, en beslutslogg, och att aldrig investera mer än du sover gott med."),
]

for rubrik_text, beskrivning in risker:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"{rubrik_text}:  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(beskrivning)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = BLACK

rubrik2("Målet är tillväxt — med öppna ögon")
brödtext(
    "Syftet med Strategiportföljen är inte att undvika risk. "
    "Det är att ta välkalibrerad risk — mer risk på det du förstår och tror på, "
    "mindre risk på det du inte gör. Ankaret ger dig marknadsuppgång med låg risk. "
    "Berättelserna ger dig chansen till home runs med liten insats. "
    "Däremellan bygger du en portfölj som kan stå upp i storm."
)

citat(
    "The stock market is a device for transferring money from the impatient to the patient."
    " — Warren Buffett"
)

rubrik3("Nödutgångar — när allt annat sviker")
brödtext(
    "Utöver MA200-regeln beräknar appen nödutgångar: om en aktie faller till 90 % av "
    "ditt genomsnittliga anskaffningsvärde (GAV), visas en signal. "
    "För kategorier 3–6 är detta en hård stopp — tid att omvärdera eller sälja. "
    "För kategorier 1–2 är det en mjuk varning — tid att köpa mer, om du tror på bolaget."
)
brödtext(
    "Disciplin i nedsida är lika viktigt som disciplin i uppsida. "
    "Den som inte vet när de ska sälja, vet egentligen inte varför de köpte."
)

sidbyte()


# ════════════════════════════════════════════════════════════
#  SIDA 5 — APPENS FUNKTIONER + KOMMA IGÅNG
# ════════════════════════════════════════════════════════════
rubrik1("5. Appens funktioner i korthet")
rubrik2("v2.08: Åtta testronder avklarade — nu kör vi skarpt")

funktioner = [
    ("Dashboard",
     "Portföljöversikt med totalt värde, nyckeltal, FX-justerade signaler och "
     "kategorifördelning. Nyckeltal uppdateras dynamiskt för vald period (30D / 90D / 6M / 1Å / Allt)."),
    ("Innehav",
     "Alla positioner med MA200-avstånd, bolagsvinst vs. valutavinst, "
     "tvådagarsstatus och nödutgångssignal. Redigerbar inline."),
    ("Transaktioner",
     "Komplett historik från Avanza. Utdelningar kopplas automatiskt till rätt aktie via ISIN."),
    ("Beslutslogg",
     "Veckovisa anteckningar om marknad och beslut. Exporterbar som CSV och Excel."),
    ("Diagram",
     "Värdeutveckling med periodväljare och referenslinje för nettoinsatt kapital. "
     "Byggs automatiskt från importerade positionsfiler."),
    ("Kassa",
     "Tillgängligt för köp per konto (ISK, AF, Sparkonto). "
     "Manuellt angivet för exakthet — Avanza inkluderar osettlerade affärer som inte syns i CSV."),
    ("Kategorier",
     "Lägg till, redigera och ta bort kategorier med namn, färg, emoji, "
     "MA200-regler och målvikter. Allt sparas lokalt."),
    ("Backup & Återställning",
     "Exportera all portföljdata till ett Excel-ark med 7 flikar. "
     "Importera tillbaka för fullständig återställning — t.ex. vid byte av enhet."),
]

for rubrik_text, beskrivning in funktioner:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"▸  {rubrik_text}  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = NAVY
    r2 = p.add_run("— " + beskrivning)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = BLACK

linje()

rubrik1("6. Kom igång på fem minuter")
rubrik2("v2.08: Åtta ark exporteras — men det räcker med tre CSV-filer för att börja")

steg = [
    ("Öppna appen",
     "Gå till https://stockholmsvy1-droid.github.io/strategiportfoljen/ i Safari (iPad) "
     "eller valfri webbläsare på dator. Bokmärk sidan — du återkommer varje vecka."),
    ("Exportera från Avanza",
     "Logga in på Avanza i webbläsaren (inte appen). Gå till:\n"
     "  • Min ekonomi → Innehav → Exportera → positioner_DATUM.csv\n"
     "  • Min ekonomi → Transaktioner → Exportera transaktioner → transaktioner_DATUM.csv\n"
     "  • (Första gången) Min ekonomi → Innehav → Exportera inköpskurser → inkopskurs_DATUM.csv"),
    ("Importera filerna",
     "Gå till fliken Importera. Dra eller välj de tre CSV-filerna i rätt kort. "
     "Appen läser dem och bygger upp innehav, transaktioner och historik automatiskt."),
    ("Sätt MA200-värden",
     "I Innehav-fliken — klicka på ett innehav och ange MA200 i lokal valuta. "
     "Hitta värdet via din mäklare eller finance.yahoo.com. "
     "Du behöver bara göra detta en gång per innehav."),
    ("Anpassa kategorier (valfritt)",
     "Gå till Kategorier-fliken och byt namn, färg eller MA200-regler till det som passar din strategi."),
    ("Exportera backup",
     "Gå till Importera → Backup & Återställning → Ladda ner Excel-backup. "
     "Spara filen i iCloud eller OneDrive. Gör detta varje vecka."),
]

for i, (rubrik_text, beskrivning) in enumerate(steg, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"Steg {i}: {rubrik_text}")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY

    lines = beskrivning.split("\n")
    for line in lines:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.8)
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(line.strip())
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = BLACK

linje()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Strategiportföljen v2.08  ·  Byggt för Martin  ·  Strategi från januari 2026")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = MUTED


# ════════════════════════════════════════════════════════════
#  Spara
# ════════════════════════════════════════════════════════════
utfil = r"C:\Users\hejma\Projekt_Claude\strategiportfoljen\Strategiportfoljen_Beskrivning_v208.docx"
doc.save(utfil)
print(f"Skapad: {utfil}")
